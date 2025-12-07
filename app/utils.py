import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple

#  - 每个样式只有一个捕获组是页码数字；下面的 infer_page_map 会找出命中的那个分组
_PAGE_PATTERNS = [
    r"(?:^|\n)\s*Page\s*(\d+)\s*(?:\n|$)",            # Page 12
    r"(?:^|\n)\s*p\.\s*(\d+)\s*(?:\n|$)",             # p. 12
    r"(?:^|\n)\s*頁\s*(\d+)\s*(?:\n|$)",              # 頁 12
    r"(?:^|\n)\s*第\s*(\d+)\s*頁\s*(?:\n|$)",          # 第 12 頁
    r"(?:^|\n)\s*Page\s*(\d+)\s*of\s*\d+\s*(?:\n|$)", # Page 12 of 200
    # —— HTML / PDF 转 Markdown 常见锚点 ——
    r"<span[^>]*\bid=['\"]page-(\d+)(?:-[^'\"]+)?['\"][^>]*>",   # <span id="page-30-0">
    r"<a[^>]*\b(?:id|name)=['\"]page-(\d+)['\"][^>]*>",          # <a id="page-30"> / <a name="page-30">
    r"<div[^>]*\bclass=['\"][^'\"]*\bpage\b[^'\"]*['\"][^>]*\bdata-page=['\"](\d+)['\"][^>]*>",  # <div class="page" data-page="30">
]
_PAGE_RE = re.compile("|".join(f"(?:{p})" for p in _PAGE_PATTERNS), re.IGNORECASE)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {'.md', '.markdown', '.pdf', '.txt', '.docx'}

def read_pdf_file(file_path: Path) -> Tuple[str, List[Dict]]:
    """
    读取PDF文件，返回文本内容和按页分割的信息。
    返回: (full_text, page_ranges)
    page_ranges: [{"page": 1, "start": 0, "end": 500}, ...]
    """
    try:
        import pdfplumber
    except ImportError:
        print("[WARN] pdfplumber not installed. Run: pip install pdfplumber")
        return "", []

    full_text = ""
    page_ranges = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                start_pos = len(full_text)

                # 提取页面文本
                page_text = page.extract_text() or ""

                # 尝试提取表格（转为文本格式）
                tables = page.extract_tables()
                table_text = ""
                for table in tables:
                    if table:
                        for row in table:
                            row_text = " | ".join(str(cell) if cell else "" for cell in row)
                            table_text += row_text + "\n"

                # 合并页面文本和表格文本
                combined_text = page_text
                if table_text and table_text.strip() not in page_text:
                    combined_text += "\n\n[Table Content]\n" + table_text

                # 添加页码标记便于后续定位
                full_text += f"\n\n--- Page {i} ---\n\n{combined_text}"

                end_pos = len(full_text)
                page_ranges.append({"page": i, "start": start_pos, "end": end_pos})

    except Exception as e:
        print(f"[WARN] Failed to read PDF {file_path}: {e}")
        return "", []

    return full_text.strip(), page_ranges

def read_docx_file(file_path: Path) -> str:
    """读取DOCX Word文档"""
    try:
        from docx import Document
    except ImportError:
        print("[WARN] python-docx not installed. Run: pip install python-docx")
        return ""

    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格内容
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                paragraphs.append("\n[Table]\n" + "\n".join(table_text))

        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"[WARN] Failed to read DOCX {file_path}: {e}")
        return ""

def read_txt_file(file_path: Path) -> str:
    """读取TXT纯文本文件，自动检测编码"""
    try:
        import chardet
    except ImportError:
        chardet = None

    # 首先尝试UTF-8
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        pass

    # 使用chardet检测编码
    if chardet:
        try:
            raw = file_path.read_bytes()
            detected = chardet.detect(raw)
            encoding = detected.get('encoding', 'utf-8')
            return raw.decode(encoding, errors='ignore')
        except Exception:
            pass

    # 最后尝试常见编码
    for enc in ['gbk', 'gb2312', 'big5', 'latin-1']:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue

    return file_path.read_text(encoding="utf-8", errors="ignore")

def read_document(file_path: Path) -> Dict:
    """
    统一的文档读取接口，支持多种格式。
    返回: {"path": str, "text": str, "format": str, "page_ranges": list (可选)}
    """
    suffix = file_path.suffix.lower()

    if suffix in {'.md', '.markdown'}:
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return {"path": str(file_path), "text": text, "format": "markdown"}
        except Exception as e:
            print(f"[WARN] Failed to read {file_path}: {e}")
            return {"path": str(file_path), "text": "", "format": "markdown"}

    elif suffix == '.pdf':
        text, page_ranges = read_pdf_file(file_path)
        return {
            "path": str(file_path),
            "text": text,
            "format": "pdf",
            "page_ranges": page_ranges
        }

    elif suffix == '.docx':
        text = read_docx_file(file_path)
        return {"path": str(file_path), "text": text, "format": "docx"}

    elif suffix == '.txt':
        text = read_txt_file(file_path)
        return {"path": str(file_path), "text": text, "format": "txt"}

    else:
        print(f"[WARN] Unsupported file format: {suffix}")
        return {"path": str(file_path), "text": "", "format": "unknown"}

def read_all_documents(docs_dir: str) -> List[Dict]:
    """
    读取目录下所有支持格式的文档。
    支持: .md, .markdown, .pdf, .txt, .docx
    """
    docs_path = Path(docs_dir)
    docs = []

    # 收集所有支持的文件
    all_files = []
    for ext in SUPPORTED_EXTENSIONS:
        all_files.extend(docs_path.glob(f"**/*{ext}"))

    # 去重并排序
    all_files = sorted(set(all_files))

    for p in all_files:
        doc = read_document(p)
        if doc["text"]:  # 只保留有内容的文档
            docs.append(doc)
            print(f"[ingest] Loaded {p.name} ({doc['format']}, {len(doc['text'])} chars)")

    return docs

def read_markdown_files(docs_dir: str) -> List[Dict]:
    """
    兼容旧接口：读取所有支持的文档（不仅仅是Markdown）
    """
    return read_all_documents(docs_dir)

def infer_page_map(md_text: str) -> List[Dict]:
    matches = list(_PAGE_RE.finditer(md_text))
    if not matches:
        return [{"page": None, "start": 0, "end": len(md_text)}]

    def first_int_group(m: re.Match) -> int | None:
        # 由于我们是一个大 OR，groups() 里只有一个是数字，其它是 None
        for g in m.groups():
            if g and g.isdigit():
                return int(g)
        return None

    pages = []
    for i, m in enumerate(matches):
        page_no = first_int_group(m)
        if page_no is None:
            # 理论不会发生；防御性处理
            continue
        start = m.start()
        end = matches[i+1].start() if i+1 < len(matches) else len(md_text)
        pages.append({"page": page_no, "start": start, "end": end})
    return pages